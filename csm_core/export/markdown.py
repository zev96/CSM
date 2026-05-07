"""Export the finished article as .md or .docx — content only.

The older snapshot sidecar (``{stem}.assembly.json``) was dropped: users
asked for plain document export, not a bundle that includes the plan /
settings. Downstream tooling that needs the plan can serialize it
separately.

Filenames follow ``MMDD-N`` — month/day plus a per-day sequence number
that picks the next available slot in the output directory. So the
first export on April 27 becomes ``0427-1`` and the next ``0427-2``,
across both .md and .docx.
"""
from __future__ import annotations
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Literal
from ..assembler.plan import AssemblyPlan

ExportFormat = Literal["markdown", "docx"]


def _heading_level(line: str) -> tuple[int, str] | None:
    """Return (level, text) for a markdown heading line, else None."""
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if not m:
        return None
    return len(m.group(1)), m.group(2).strip()


# Inline markdown bold ``**text**``. Non-greedy, single-line, requires at
# least one inner char. Used by the docx writer to apply Word's bold run
# property instead of leaving the raw asterisks visible.
_BOLD_RE = re.compile(r"\*\*([^*\n]+?)\*\*")


def _strip_bold_markers(text: str) -> str:
    """Drop ``**...**`` wrappers but keep the inner text.

    Used for heading bodies — Word heading styles are already bold, so
    the asterisks are pure noise.
    """
    return _BOLD_RE.sub(r"\1", text)


def extract_title(text: str) -> str:
    """Pull the first H1 / H2 line from a markdown document.

    Used by the home-page recents to show the article's headline rather
    than the on-disk filename. Falls back to the first non-empty prose
    line if no headings exist; returns "" if the document is empty.
    """
    first_prose = ""
    for raw in (text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        h = _heading_level(line)
        if h is not None and h[1]:
            return h[1]
        if not first_prose:
            first_prose = line
    return first_prose


def _next_filename_stem(out_dir: Path, fmt: ExportFormat,
                       now: datetime | None = None) -> str:
    """Return ``MMDD-N`` where N is the smallest free index for today.

    Looks at every ``MMDD-*.md`` and ``MMDD-*.docx`` already in
    ``out_dir`` and picks the next integer so re-exporting the same day
    doesn't clobber existing files.
    """
    now = now or datetime.now()
    prefix = now.strftime("%m%d")
    pattern = re.compile(rf"^{prefix}-(\d+)\.(md|docx)$", re.IGNORECASE)
    used: set[int] = set()
    if out_dir.exists():
        for p in out_dir.iterdir():
            m = pattern.match(p.name)
            if m:
                used.add(int(m.group(1)))
    n = 1
    while n in used:
        n += 1
    return f"{prefix}-{n}"


# ── Docx rendering ──────────────────────────────────────────────────────────
_DOCX_FONT = "Microsoft YaHei"  # 微软雅黑
_HEADING_COLOR = (0x1E, 0x1C, 0x19)  # near-black, matches the app's _INK token
_LINE_SPACING = 1.5


def _set_run_font(run, *, color: tuple[int, int, int] | None = None) -> None:
    """Force both Western and East-Asian font slots to 微软雅黑 + colour.

    python-docx only sets the Western font via ``run.font.name``; Word
    falls back to the theme East-Asian font for CJK runs unless we also
    write the ``w:eastAsia`` attribute on ``w:rFonts``. ``color`` is
    forced explicitly because the built-in Heading styles ship with a
    blue accent that overrides plain text colour.
    """
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    run.font.name = _DOCX_FONT
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), _DOCX_FONT)


def _apply_default_font(doc) -> None:
    """Set Normal style to 微软雅黑 + 1.5 line spacing + black text."""
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    style = doc.styles["Normal"]
    style.font.name = _DOCX_FONT
    style.font.color.rgb = RGBColor(*_HEADING_COLOR)
    pf = style.paragraph_format
    pf.line_spacing = _LINE_SPACING
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), _DOCX_FONT)


def _override_heading_styles(doc) -> None:
    """Repaint the built-in Heading 1-4 styles to black + 微软雅黑.

    Word ships these styles with a blue accent that survives most run-
    level overrides because Word resolves character properties bottom-
    up. Updating the style itself avoids the colour creeping back when
    Word re-applies theme colours on open.
    """
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    for level in (1, 2, 3, 4):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        style.font.name = _DOCX_FONT
        style.font.color.rgb = RGBColor(*_HEADING_COLOR)
        pf = style.paragraph_format
        pf.line_spacing = _LINE_SPACING
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), _DOCX_FONT)


def _add_inline_runs(paragraph, text: str) -> None:
    """Append runs to *paragraph*, parsing inline ``**bold**`` markup.

    Splits *text* into alternating plain / bold segments and emits each
    as a separate run with the right ``run.bold`` flag. The Word run
    font is also reset on every run so the document's default 黑色 +
    OPPOSans styling stays consistent.
    """
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_run_font(run, color=_HEADING_COLOR)
        run = paragraph.add_run(m.group(1))
        run.bold = True
        _set_run_font(run, color=_HEADING_COLOR)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, color=_HEADING_COLOR)


def _write_docx(path: Path, text: str) -> None:
    """Render the markdown as a .docx — headings + paragraphs.

    Paragraph grouping follows markdown semantics: blocks are separated
    by blank lines. Headings (`#`…`######`) become Word heading styles
    repainted black; everything else becomes a Normal paragraph with
    inline ``**bold**`` markup applied via Word's run-bold property.
    All paragraphs run at 1.5 line spacing.
    """
    from docx import Document
    doc = Document()
    _apply_default_font(doc)
    _override_heading_styles(doc)

    # Split into blocks on blank-line boundaries so multi-line markdown
    # paragraphs are emitted as a single Word paragraph.
    blocks: list[list[str]] = []
    current: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)
    if current:
        blocks.append(current)

    for block in blocks:
        first = block[0]
        heading = _heading_level(first)
        if heading is not None and len(block) == 1:
            level, body = heading
            # Word heading 样式自带粗体，去掉 ``**...**`` 包裹只保留文字
            # 否则会看到字面星号。
            body = _strip_bold_markers(body)
            p = doc.add_heading(level=min(level, 4))
            p.paragraph_format.line_spacing = _LINE_SPACING
            run = p.add_run(body)
            _set_run_font(run, color=_HEADING_COLOR)
            continue
        # Treat block as a single paragraph. Line breaks inside a block
        # are preserved as soft returns so the user can see the original
        # layout but the text still flows as one paragraph.
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = _LINE_SPACING
        for i, line in enumerate(block):
            if i > 0:
                # Soft line break inside the same paragraph.
                br_run = p.add_run()
                br_run.add_break()
                _set_run_font(br_run, color=_HEADING_COLOR)
            _add_inline_runs(p, line)

    doc.save(str(path))


def export_article(
    *,
    out_dir: Path,
    keyword: str,                                # kept for call-site compat
    final_text: str,
    plan: AssemblyPlan | None = None,            # kept for call-site compat
    prompt_snapshot: dict[str, Any] | None = None,  # unused now
    fmt: ExportFormat = "markdown",
) -> dict[str, str]:
    out_dir = Path(out_dir)
    if not out_dir.exists():
        raise FileNotFoundError(f"output directory does not exist: {out_dir}")

    stem = _next_filename_stem(out_dir, fmt)
    title = extract_title(final_text) or keyword

    if fmt == "docx":
        path = out_dir / f"{stem}.docx"
        _write_docx(path, final_text)
        return {"document": str(path), "format": "docx", "title": title}

    path = out_dir / f"{stem}.md"
    path.write_text(final_text, encoding="utf-8")
    return {"document": str(path), "format": "markdown", "title": title}
