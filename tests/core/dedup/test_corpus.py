"""Corpus scanner: walk dir, extract text from .md/.docx/.txt, track mtime."""
from pathlib import Path
import pytest
from csm_core.dedup.corpus import (
    CorpusEntry, scan_corpus, extract_text, extract_title,
)


def test_extract_text_from_md(tmp_path: Path):
    f = tmp_path / "test.md"
    f.write_text("# Title\n\n正文内容第一段\n\n第二段", encoding="utf-8")
    text = extract_text(f)
    assert "正文内容第一段" in text
    assert "第二段" in text


def test_extract_text_from_txt(tmp_path: Path):
    f = tmp_path / "test.txt"
    f.write_text("纯文本内容", encoding="utf-8")
    assert extract_text(f) == "纯文本内容"


def test_extract_text_from_docx(tmp_path: Path):
    """Smoke: build a real .docx via python-docx and read it back."""
    from docx import Document
    doc_path = tmp_path / "test.docx"
    d = Document()
    d.add_paragraph("第一段内容")
    d.add_paragraph("第二段内容")
    d.save(str(doc_path))
    text = extract_text(doc_path)
    assert "第一段内容" in text
    assert "第二段内容" in text


def test_extract_text_unsupported_format_returns_empty(tmp_path: Path):
    f = tmp_path / "test.bin"
    f.write_bytes(b"\x00\x01\x02")
    assert extract_text(f) == ""


def test_extract_title_md_uses_first_h1(tmp_path: Path):
    f = tmp_path / "test.md"
    f.write_text("# 我的文章标题\n\n正文", encoding="utf-8")
    assert extract_title(f) == "我的文章标题"


def test_extract_title_md_falls_back_to_stem(tmp_path: Path):
    f = tmp_path / "no-h1.md"
    f.write_text("没有标题的笔记", encoding="utf-8")
    assert extract_title(f) == "no-h1"


def test_scan_corpus_yields_entries(tmp_path: Path):
    (tmp_path / "a.md").write_text("# A\n\n内容 A", encoding="utf-8")
    (tmp_path / "b.txt").write_text("内容 B", encoding="utf-8")
    (tmp_path / "ignored.bin").write_bytes(b"\x00")
    sub = tmp_path / "sub"
    sub.mkdir()
    (sub / "c.md").write_text("# C\n\n内容 C", encoding="utf-8")

    entries = list(scan_corpus(tmp_path))
    paths = {e.path.name for e in entries}
    assert paths == {"a.md", "b.txt", "c.md"}
    for e in entries:
        assert isinstance(e, CorpusEntry)
        assert e.text  # non-empty
        assert e.mtime > 0


def test_scan_corpus_missing_dir_returns_empty(tmp_path: Path):
    nonexist = tmp_path / "doesnotexist"
    assert list(scan_corpus(nonexist)) == []


def test_scan_corpus_skips_huge_file(tmp_path: Path):
    """File over 5MB is skipped (likely not a finished article)."""
    big = tmp_path / "big.md"
    big.write_text("x" * (6 * 1024 * 1024), encoding="utf-8")
    entries = list(scan_corpus(tmp_path))
    assert big not in [e.path for e in entries]
